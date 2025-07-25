import streamlit as st
import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
from docx import Document
import io

# Streamlit page configuration
st.title("PDF to Telugu Translated DOCX Converter")
st.write("Upload a PDF file, click the translate button, and download the translated content as a DOCX file.")

# File uploader for PDF
uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

if uploaded_file is not None:
    # Read the uploaded PDF file to determine page count
    pdf_bytes = uploaded_file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)
    doc.close()  # Close temporarily, reopen during translation

    # Button to start translation
    if st.button("Start Translation"):
        try:
            # Reopen PDF for processing
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")

            # Initialize the translator
            translator = GoogleTranslator(source='auto', target='te')

            # Create a new Word document
            document = Document()
            document.add_heading('Translated PDF (Telugu)', level=1)

            # Initialize progress bar
            progress_bar = st.progress(0)
            progress_text = st.empty()

            # Extract and translate text from each page
            for i, page in enumerate(doc):
                # Update progress bar
                progress = (i + 1) / total_pages
                progress_bar.progress(progress)
                progress_text.text(f"Translating page {i+1} of {total_pages}...")

                text = page.get_text()
                if text.strip():
                    try:
                        translated = translator.translate(text)
                        document.add_paragraph(f"Page {i+1}", style='Heading 2')
                        for line in translated.split('\n'):
                            if line.strip():
                                document.add_paragraph(line)
                    except Exception as e:
                        st.warning(f"Translation error on page {i+1}: {e}")

            # Finalize progress bar
            progress_bar.progress(1.0)
            progress_text.text("Translation complete!")

            doc.close()

            # Save the document to a bytes buffer
            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)

            # Provide download button for the DOCX file
            st.download_button(
                label="Download Translated DOCX",
                data=buffer,
                file_name="translated_telugu.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("Translation complete! Click the button to download the DOCX file.")

        except Exception as e:
            st.error(f"An error occurred: {e}")
            progress_bar.empty()
            progress_text.empty()
else:
    st.info("Please upload a PDF file to begin.")